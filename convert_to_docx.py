"""
仿生蝴蝶扑翼MAV机械原理分析报告 — Markdown → Word 转换脚本
两步法：pandoc(md→docx) + python-docx(格式后处理)
"""

import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# 配置
# ============================================================
PROJECT_ROOT = Path(__file__).parent
REPORT_DIR = PROJECT_ROOT / "report"
OUTPUT_DIR = PROJECT_ROOT / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

# 章节文件顺序
CHAPTER_FILES = [
    "摘要.md",
    "01_绪论.md",
    "02_整机总体构型.md",
    "03_轮系传动设计_上半部分.md",
    "03_轮系传动设计_下半部分.md",
    "04_曲柄摇杆机构_上半部分.md",
    "04_曲柄摇杆机构_下半部分.md",
    "05_丝杆调节机构.md",
    "06_气动外载分析.md",
    "07_机械动力学与载荷校核.md",
    "08_结论与展望.md",
    "心得体会.md",
]

# 格式常量
FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TNR = "Times New Roman"

PT_16 = Pt(16)   # 三号 - 章标题
PT_15 = Pt(15)   # 小三号 - 节标题
PT_14 = Pt(14)   # 四号 - 条标题
PT_12 = Pt(12)   # 小四号 - 正文
PT_10_5 = Pt(10.5)  # 五号 - 图题/表题/参考文献
PT_9 = Pt(9)     # 小五号 - 页眉页脚

LINE_SPACING_20PT = Pt(20)  # 固定值20磅
FIRST_LINE_INDENT = Cm(0.74)  # 约2字符（小四号宋体）

MAX_IMG_WIDTH_CM = 14.0  # 图片最大宽度

PAGE_WIDTH_CM = 21.0   # A4
PAGE_HEIGHT_CM = 29.7  # A4
MARGIN_TOP_CM = 2.54
MARGIN_BOTTOM_CM = 2.54
MARGIN_LEFT_CM = 3.17
MARGIN_RIGHT_CM = 3.17


# ============================================================
# Step 1: 合并 md 文件
# ============================================================
def concatenate_markdown():
    """合并所有 md 文件，修正图片路径为绝对路径"""
    combined = []

    for fname in CHAPTER_FILES:
        fpath = REPORT_DIR / fname
        if not fpath.exists():
            print(f"  WARNING: {fpath} not found, skipping")
            continue

        with open(fpath, "r", encoding="utf-8") as f:
            content = f.read()

        # 修正图片路径：相对路径 → 绝对路径
        def fix_img_path(match):
            prefix = match.group(1)  # ![alt](
            path = match.group(2)    # 路径
            suffix = match.group(3)  # )

            # 如果是相对路径（以 ../ 开头）
            if path.startswith("../"):
                # 解析相对于 report/ 的路径
                abs_path = (REPORT_DIR / path).resolve()
                return f"{prefix}{abs_path}{suffix}"
            # 如果已经是绝对路径或不存在
            return match.group(0)

        content = re.sub(
            r'(!\[[^\]]*\]\()([^)\s]+)(\))',
            fix_img_path,
            content
        )

        combined.append(content)
        combined.append("\n\n")  # 章节间空行

    combined_md = "\n".join(combined)

    # 写入临时文件
    tmp_md = PROJECT_ROOT / "temp_combined.md"
    with open(tmp_md, "w", encoding="utf-8") as f:
        f.write(combined_md)

    print(f"  Combined markdown: {tmp_md} ({len(combined_md)} chars)")
    return tmp_md


# ============================================================
# Step 2: pandoc 转换
# ============================================================
def run_pandoc(md_path):
    """使用 pandoc 将 md 转换为 docx"""
    output_docx = OUTPUT_DIR / "report_raw.docx"

    cmd = [
        "pandoc",
        str(md_path),
        "-o", str(output_docx),
        "--from=markdown+tex_math_dollars",
        "--to=docx",
        "--resource-path=" + str(PROJECT_ROOT),
        "--wrap=none",
    ]

    print(f"  Running: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True)

    if result.returncode != 0:
        print(f"  pandoc ERROR: {result.stderr}")
        sys.exit(1)

    print(f"  pandoc output: {output_docx}")
    return output_docx


# ============================================================
# Step 3: python-docx 格式后处理
# ============================================================

def set_cell_font(cell, font_name=FONT_SONG, font_size=PT_10_5, bold=False):
    """设置单元格字体"""
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = bold
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def format_table_三線表(table):
    """将表格格式化为三线表"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 清除所有边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # 设置表头行下边框（栏目线）
    if len(table.rows) > 1:
        header_row = table.rows[1]  # 第一行数据行（跳过可能的标题行）
        for cell in header_row.cells:
            tc = cell._tc
            tcPr = tc.tcPr if tc.tcPr is not None else parse_xml(f'<w:tcPr {nsdecls("w")}/>')
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                '</w:tcBorders>'
            )
            tcPr.append(borders)

    # 设置单元格字体和居中
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = FONT_TNR
                    run.font.size = PT_10_5
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)


def format_document(docx_path):
    """应用所有格式要求"""
    doc = Document(docx_path)

    # --- 页面设置 ---
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)

    # --- 收集所有章节标题，用于页眉右侧 ---
    chapter_titles = {}
    current_chapter = "摘要"
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            current_chapter = para.text.strip()
        elif para.style.name == 'Heading 2':
            chapter_titles[para.text.strip()] = current_chapter

    # --- 页眉：左侧"机械原理大作业" + 右侧当前章标题 ---
    header = section.header
    header_para = header.paragraphs[0]
    header_para.clear()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 左侧文字
    left_run = header_para.add_run("机械原理大作业")
    left_run.font.name = FONT_SONG
    left_run.font.size = PT_9
    left_run.font.color.rgb = RGBColor(0, 0, 0)
    left_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)

    # 中间点线填充（用 Tab 实现）
    tab_run = header_para.add_run("\t")
    tab_run.font.name = FONT_SONG
    tab_run.font.size = PT_9
    tab_run.font.color.rgb = RGBColor(0, 0, 0)
    tab_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)

    # 右侧章标题（用 StyleRef 字段引用当前章标题）
    right_run = header_para.add_run()
    right_run.font.name = FONT_SONG
    right_run.font.size = PT_9
    right_run.font.color.rgb = RGBColor(0, 0, 0)
    right_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)
    # 添加 StyleRef 字段引用 Heading 1
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> STYLEREF 1 \\* MERGEFORMAT </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    right_run._element.append(fldChar1)
    right_run._element.append(instrText)
    right_run._element.append(fldChar2)

    # --- 页脚（页码）---
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run()
    footer_run.font.name = FONT_TNR
    footer_run.font.size = PT_9
    footer_run.font.color.rgb = RGBColor(0, 0, 0)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    footer_run._element.append(fldChar1)
    footer_run._element.append(instrText)
    footer_run._element.append(fldChar2)

    # --- 辅助函数：确保 run 为黑色 ---
    def ensure_black(run):
        run.font.color.rgb = RGBColor(0, 0, 0)

    # --- 遍历段落应用格式 ---
    for para in doc.paragraphs:
        text = para.text.strip()

        # 章标题 (H1): # 开头
        if para.style.name == 'Heading 1':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(24)
            para.paragraph_format.space_after = Pt(18)
            for run in para.runs:
                run.font.name = FONT_HEI
                run.font.size = PT_16
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEI)

        # 节标题 (H2): ## 开头
        elif para.style.name == 'Heading 2':
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.name = FONT_HEI
                run.font.size = PT_15
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEI)

        # 条标题 (H3): ### 开头
        elif para.style.name == 'Heading 3':
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(3)
            for run in para.runs:
                run.font.name = FONT_HEI
                run.font.size = PT_14
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEI)

        # 图题/表题段落（以"图X-X"或"表X-X"开头）
        elif re.match(r'^[图表]\d+-\d+', text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.name = FONT_SONG
                run.font.size = PT_10_5
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)

        # 子图标注（以"（a）"或"（b）"开头）
        elif re.match(r'^[（(][ab][）)]', text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = FONT_SONG
                run.font.size = PT_10_5
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)

        # 公式行（以 $$ 开头或包含公式编号）
        elif text.startswith('$$') or re.search(r'\(\d+-\d+\)', text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)

        # 正文段落
        else:
            if text:  # 非空段落
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
                para.paragraph_format.line_spacing = LINE_SPACING_20PT

                for run in para.runs:
                    run.font.name = FONT_SONG
                    run.font.size = PT_12
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)

                    # 英文/数字用 Times New Roman
                    if re.search(r'[a-zA-Z0-9]', run.text):
                        run.font.name = FONT_TNR
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONG)

    # --- 格式化表格 ---
    for table in doc.tables:
        format_table_三線表(table)
        # 确保表格内所有文字为黑色
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)

    # --- 处理图片 ---
    for para in doc.paragraphs:
        for run in para.runs:
            ensure_black(run)

        # 检查段落中的 inline 图片
        for inline in para._element.findall(f'.//{qn("wp:inline")}'):
            extent = inline.find(qn('wp:extent'))
            if extent is not None:
                cx = int(extent.get('cx', 0))
                cy = int(extent.get('cy', 0))

                width_cm = cx / 360000
                height_cm = cy / 360000

                if width_cm > MAX_IMG_WIDTH_CM:
                    scale = MAX_IMG_WIDTH_CM / width_cm
                    new_cx = int(cx * scale)
                    new_cy = int(cy * scale)
                    extent.set('cx', str(new_cx))
                    extent.set('cy', str(new_cy))

    # 保存
    output_path = OUTPUT_DIR / "仿生蝴蝶扑翼MAV机械原理分析报告.docx"
    doc.save(output_path)
    print(f"  Saved: {output_path}")
    return output_path


# ============================================================
# 主流程
# ============================================================
def main():
    print("=" * 60)
    print("仿生蝴蝶扑翼MAV机械原理分析报告 — 格式转换")
    print("=" * 60)

    # Step 1: 合并 md
    print("\n[Step 1] 合并 Markdown 文件...")
    md_path = concatenate_markdown()

    # Step 2: pandoc 转换
    print("\n[Step 2] pandoc 转换 md → docx...")
    raw_docx = run_pandoc(md_path)

    # Step 3: 格式后处理
    print("\n[Step 3] python-docx 格式后处理...")
    final_docx = format_document(raw_docx)

    print("\n" + "=" * 60)
    print(f"完成！输出文件：{final_docx}")
    print("=" * 60)


if __name__ == "__main__":
    main()
