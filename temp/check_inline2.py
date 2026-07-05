from docx import Document
from docx.oxml.ns import qn

doc = Document('output/仿生蝴蝶扑翼MAV机械原理分析报告.docx')

# Search for any paragraph with inline math
found = False
for para in doc.paragraphs:
    # Check if paragraph has runs with math elements
    for run in para.runs:
        math_elems = run._element.findall(qn('m:oMath'))
        if math_elems:
            print(f'找到行内公式段落：{para.text[:80]}')
            print(f'  Run 文本：{run.text[:50] if run.text else "(empty)"}')
            print(f'  公式数：{len(math_elems)}')
            found = True
            break
    if found:
        break

if not found:
    print('未找到行内公式')
    # Let's check what pandoc produces for inline math
    print('\n检查包含 $ 的段落...')
    for para in doc.paragraphs:
        if '$' in para.text:
            print(f'段落：{para.text[:100]}')
            break
