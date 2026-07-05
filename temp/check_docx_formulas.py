"""检查 docx 中的公式"""
from docx import Document
from docx.oxml.ns import qn

doc = Document('output/仿生蝴蝶扑翼MAV机械原理分析报告.docx')

# Count equations (OMML math elements)
equation_count = 0
for para in doc.paragraphs:
    # Check for math elements
    math_elements = para._element.findall(qn('m:oMath'))
    if math_elements:
        equation_count += len(math_elements)
        # Get formula text
        for math in math_elements:
            # Get all text in the math element
            texts = []
            for t in math.iter(qn('m:t')):
                if t.text:
                    texts.append(t.text)
            formula = ''.join(texts)
            if formula:
                print(f'公式：{formula[:80]}')

print(f'\n总公式数：{equation_count}')

# Check for $$ that might not have been converted
print('\n检查未转换的 $$...')
for i, para in enumerate(doc.paragraphs):
    if '$$' in para.text:
        print(f'段落 {i}: {para.text[:100]}')
